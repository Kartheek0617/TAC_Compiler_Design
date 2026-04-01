import sys
from docx import Document
from docx.shared import Pt, Inches

def create_word_doc(input_file, output_file):
    document = Document()
    
    # Title
    document.add_heading('Project Report: Full-Stack Mini-Compiler Web IDE', 0)
    
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                document.add_paragraph('--- Code / Structure ---', style='Intense Quote')
            continue
            
        if in_code_block:
            p = document.add_paragraph(line)
            p.style = document.styles['No Spacing']
            continue
            
        if line.startswith('# '):
            continue # Skip title, already added
        elif line.startswith('## '):
            document.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=2)
        elif line.startswith('* '):
            document.add_paragraph(line[2:], style='List Bullet')
        else:
            document.add_paragraph(line)
            
    document.save(output_file)
    print(f"Successfully saved {output_file}")

if __name__ == '__main__':
    create_word_doc('Project_Report.md', 'Project_Report.docx')
